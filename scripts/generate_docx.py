import os
import io
from docx import Document
from docx.shared import Inches

try:
    import cairosvg
except Exception:
    raise SystemExit("cairosvg is required. Install with: pip install cairosvg")

BASE = os.path.dirname(os.path.dirname(__file__))
IMAGE_DIR = os.path.join(BASE, 'public')
OUT_DIR = os.path.join(BASE, 'output')
os.makedirs(OUT_DIR, exist_ok=True)

images = [
    'window.svg',
    'vercel.svg',
    'next.svg',
    'globe.svg',
    'file.svg',
]

doc = Document()
doc.add_heading('Repository Diagrams', level=1)

for img in images:
    path = os.path.join(IMAGE_DIR, img)
    if not os.path.exists(path):
        doc.add_paragraph(f'Missing: {img}')
        continue
    doc.add_paragraph(img)
    if img.lower().endswith('.svg'):
        png_bytes = cairosvg.svg2png(url=path)
        bio = io.BytesIO(png_bytes)
        bio.seek(0)
        doc.add_picture(bio, width=Inches(5))
    else:
        doc.add_picture(path, width=Inches(5))

out_path = os.path.join(OUT_DIR, 'doc_with_diagrams.docx')
doc.save(out_path)
print('Wrote', out_path)
